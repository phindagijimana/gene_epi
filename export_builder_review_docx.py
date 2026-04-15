#!/usr/bin/env python3
"""Convert gene_br.md (basic structure) to build_reviewer.docx for local reviewer handoff."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_runs_with_bold(p, text: str) -> None:
    """Split on **bold** fragments."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)


def main() -> None:
    root = Path(__file__).resolve().parent
    md_path = root / "gene_br.md"
    out_path = root / "build_reviewer.docx"
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    doc.core_properties.title = "Builder Review — Gene-Epi ILAE MAGMA pipeline"
    doc.core_properties.subject = "Inzira Labs Builder Review style (concise)"

    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(11)

    for line in lines:
        if line.strip() == "---":
            continue
        if not line.strip():
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, line[2:].strip())
            continue
        p = doc.add_paragraph()
        add_runs_with_bold(p, line.strip())

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
